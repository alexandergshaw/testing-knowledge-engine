"""Project zip -> course materials: a PPTX lecture per unit, a Word (.docx)
LMS intro per unit, Word assignment instructions per unit, and one
deterministic machine-readable rubric CSV for the whole course. Pure
extraction + templating — no LLM.

Built around the structure the Copilot project prompt produces:

    assignments/<unit>/INSTRUCTIONS.md   # H1 = "assignment1: Topic, Names"
    assignments/<unit>/student_code.py   # teaching banner + docstring tasks
    assignments/<unit>/test_assignment.py# unittest contract = grading criteria
"""

import ast
import csv
import io
import re
import zipfile
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt as DocxPt

from .slides import (
    add_bullet_slide,
    add_code_box,
    add_content_slide,
    add_text_box,
    add_title_slide,
    deck_bytes,
    new_deck,
)

MAX_ZIP_ENTRIES = 2000
MAX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
SKIP_PARTS = {"__pycache__", "node_modules", ".git", ".next", "app", "solutions"}

# Slide content limits — lectures must stay readable.
MAX_CONCEPT_SLIDES = 6
MAX_EXPLANATION_CHARS = 700
MAX_EXAMPLE_LINES = 12

PLACEHOLDER_LITERALS = ["Your Name", "TODO", "FIXME"]


class MaterialsError(ValueError):
    """User-facing problem with the uploaded project."""


@dataclass
class Unit:
    slug: str
    path: str                 # repo-relative folder, e.g. assignments/assignment1
    title: str
    topic: str
    week: int = 0
    sections: dict = field(default_factory=dict)   # "## Heading" -> text
    concepts: list = field(default_factory=list)   # {name, explanation, example}
    tasks: list = field(default_factory=list)
    rules: list = field(default_factory=list)
    steps: list = field(default_factory=list)
    tests: list = field(default_factory=list)      # test method names
    test_file: str = ""
    code_file: str = ""
    placeholder_lines: list = field(default_factory=list)


# ---------------------------------------------------------------------------
# Zip reading


def _read_zip(zip_bytes):
    """path -> text for every safe text entry, with zip-bomb guards and the
    top-level wrapper folder (github's `repo-main/`) stripped."""
    archive = zipfile.ZipFile(io.BytesIO(zip_bytes))
    infos = [i for i in archive.infolist() if not i.is_dir()]
    if len(infos) > MAX_ZIP_ENTRIES:
        raise MaterialsError("Zip has too many files.")
    if sum(i.file_size for i in infos) > MAX_UNCOMPRESSED_BYTES:
        raise MaterialsError("Zip is too large when uncompressed.")

    files = {}
    for info in infos:
        path = info.filename.replace("\\", "/")
        parts = path.split("/")
        if ".." in parts or any(p in SKIP_PARTS for p in parts) or parts[-1].startswith("."):
            continue
        files[path] = archive.read(info).decode("utf-8", errors="replace")

    # Strip a single common root folder if present.
    roots = {p.split("/")[0] for p in files}
    if len(roots) == 1 and all("/" in p for p in files):
        files = {p.split("/", 1)[1]: text for p, text in files.items()}
    return files


# ---------------------------------------------------------------------------
# Unit parsing


def _unit_sort_key(slug):
    """assignment0..N first (numeric), then reviewN/examN paired by number
    (review before its exam), final last. Deterministic for any project."""
    match = re.match(r"([a-zA-Z]+)(\d*)$", slug)
    name = match.group(1).lower() if match else slug.lower()
    number = int(match.group(2)) if match and match.group(2) else 0
    if name == "assignment":
        return (0, number, 0, slug)
    if name == "review":
        return (1, number, 0, slug)
    if name == "exam" or name == "test":
        return (1, number, 1, slug)
    if name == "final":
        return (2, number, 0, slug)
    return (0, number, 1, slug)  # unknown unit kinds ride along with assignments


def _parse_instructions(text):
    """(title, topic, sections). Title is the H1; topic is what follows the
    colon in it ('assignment1: Variables, I/O' -> 'Variables, I/O')."""
    title, topic = "", ""
    sections, current, buffer = {}, None, []
    for line in text.splitlines():
        h1 = re.match(r"^#\s+(.+)$", line)
        h2 = re.match(r"^##\s+(.+)$", line)
        if h1 and not h2 and not title:
            title = h1.group(1).strip()
            continue
        if h2:
            if current:
                sections[current] = "\n".join(buffer).strip()
            current, buffer = h2.group(1).strip(), []
            continue
        if current is not None:
            buffer.append(line)
    if current:
        sections[current] = "\n".join(buffer).strip()
    if ":" in title:
        topic = title.split(":", 1)[1].strip()
    return title, topic or title, sections


# Concept names start with an ALL-CAPS word but may carry lowercase asides:
# "3. BRANCHING (if / elif / else) — ...". The separator is an em/en dash.
_CONCEPT_START = re.compile(r"^#\s{1,4}(\d+)\.\s+([A-Z][A-Z0-9'‐-]+[^—–]{0,60}?)\s*[—–]\s+(.*)$")


def _concept_case(name):
    """BRANCHING (if / elif / else) -> Branching (if / elif / else)."""
    return " ".join(
        word.capitalize() if word.isupper() and len(word) > 1 else word
        for word in name.split()
    )


def _parse_banner_concepts(code_text):
    """Numbered concepts from the teaching comment banner at the top of the
    starter file: '#  1. VARIABLES — explanation' followed by prose and
    deeper-indented code-example lines."""
    concepts = []
    current = None
    for line in code_text.splitlines():
        if not line.lstrip().startswith("#"):
            if line.strip():
                break  # banner over at first real code
            continue
        body = line.lstrip()[1:]
        start = _CONCEPT_START.match(line)
        if start:
            current = {
                "name": _concept_case(start.group(2).strip()),
                "explanation": start.group(3).strip(),
                "example": [],
            }
            concepts.append(current)
            continue
        if current is None:
            continue
        stripped = body.rstrip()
        if not stripped.strip("# ─-—–"):
            continue
        if stripped.startswith("      "):  # deep indent = code example line
            current["example"].append(stripped.strip())
        else:
            current["explanation"] = (
                current["explanation"] + " " + stripped.strip()
            ).strip()
    for concept in concepts:
        concept["explanation"] = re.sub(r"\s+", " ", concept["explanation"])[
            :MAX_EXPLANATION_CHARS
        ]
        concept["example"] = concept["example"][:MAX_EXAMPLE_LINES]
    return concepts


def _parse_docstring_tasks(code_text):
    """('Your job:' numbered steps, 'Remember the rules:' bullets) from any
    function docstring in the starter file."""
    tasks, rules = [], []
    try:
        tree = ast.parse(code_text)
    except SyntaxError:
        return tasks, rules
    for node in ast.walk(tree):
        if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            continue
        doc = ast.get_docstring(node) or ""
        mode = None
        for line in doc.splitlines():
            stripped = line.strip()
            if re.match(r"your job\s*:", stripped, re.IGNORECASE):
                mode = "tasks"
                continue
            if re.match(r"remember the rules\s*:", stripped, re.IGNORECASE):
                mode = "rules"
                continue
            numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
            bullet = re.match(r"^[-*]\s+(.*)$", stripped)
            if mode == "tasks" and numbered:
                tasks.append(numbered.group(2).strip())
            elif mode == "tasks" and tasks and stripped and not bullet:
                tasks[-1] = f"{tasks[-1]} {stripped}"
            elif mode == "rules" and bullet:
                rules.append(bullet.group(1).strip())
            elif stripped == "":
                continue
    return tasks, rules


_STEP_MARKER = re.compile(r"#\s*[─—–-]*\s*Step\s*(\d+)\s*:\s*(.*?)\s*[─—–-]*\s*$")


def _parse_step_markers(code_text):
    steps = []
    for line in code_text.splitlines():
        match = _STEP_MARKER.search(line)
        if match:
            steps.append(match.group(2).strip())
    return steps


def _parse_tests(test_text):
    try:
        tree = ast.parse(test_text)
    except SyntaxError:
        return []
    return sorted(
        node.name
        for node in ast.walk(tree)
        if isinstance(node, ast.FunctionDef) and node.name.startswith("test_")
    )


def _find_placeholders(code_text):
    """Exact starter-code lines a finished submission must have changed —
    deterministic sentinels for the rubric's static check."""
    lines = []
    for line in code_text.splitlines():
        stripped = line.strip()
        if any(literal in stripped for literal in PLACEHOLDER_LITERALS) and (
            "=" in stripped and not stripped.startswith("#")
        ):
            lines.append(stripped)
        elif re.match(r"^\w+\s*=\s*(\[\]|\{\}|\"\"|'')\s*(#.*)?$", stripped):
            lines.append(stripped)
    return sorted(set(lines))


def humanize_test(name):
    return name.removeprefix("test_").replace("_", " ").strip()


def parse_project(zip_bytes):
    """All teachable units in the uploaded project, in course order."""
    try:
        files = _read_zip(zip_bytes)
    except zipfile.BadZipFile:
        raise MaterialsError("That file isn't a valid zip archive.")

    folders = {}
    for path in files:
        if path.lower().endswith("instructions.md"):
            folders[path.rsplit("/", 1)[0]] = path
    if not folders:
        raise MaterialsError(
            "No units found: expected folders containing an INSTRUCTIONS.md "
            "(the structure the Copilot project prompt generates)."
        )

    units = []
    for folder, instructions_path in folders.items():
        slug = folder.rsplit("/", 1)[-1]
        title, topic, sections = _parse_instructions(files[instructions_path])
        unit = Unit(slug=slug, path=folder, title=title or slug, topic=topic or slug)
        unit.sections = sections

        in_folder = [p for p in files if p.startswith(folder + "/")]
        test_files = sorted(p for p in in_folder if re.search(r"/test_[^/]+\.py$", p))
        code_files = sorted(
            p
            for p in in_folder
            if p.endswith(".py") and p not in test_files
        )
        preferred = [p for p in code_files if p.endswith("student_code.py")]
        unit.code_file = (preferred or code_files or [""])[0]
        unit.test_file = (test_files or [""])[0]

        if unit.code_file:
            code_text = files[unit.code_file]
            unit.concepts = _parse_banner_concepts(code_text)
            unit.tasks, unit.rules = _parse_docstring_tasks(code_text)
            unit.steps = _parse_step_markers(code_text)
            unit.placeholder_lines = _find_placeholders(code_text)
        if unit.test_file:
            unit.tests = _parse_tests(files[unit.test_file])
        units.append(unit)

    units.sort(key=lambda u: _unit_sort_key(u.slug))
    for week, unit in enumerate(units, start=1):
        unit.week = week
    return units


# ---------------------------------------------------------------------------
# Generators


def build_lecture(unit):
    """One PPTX lecture for a unit, as bytes — themed, ≤2 bullets per slide."""
    deck = new_deck()
    footer = f"Week {unit.week} · {unit.topic}"
    add_title_slide(deck, f"Week {unit.week}: {unit.topic}", unit.title)

    objectives = []
    if unit.sections.get("Learning target"):
        objectives.append(unit.sections["Learning target"].splitlines()[0])
    objectives.extend(f"Understand {c['name']}" for c in unit.concepts)
    if not objectives:
        objectives = [f"Work through this week's unit: {unit.topic}"]
    add_bullet_slide(deck, "This week's goals", objectives, footer=footer)

    for concept in unit.concepts[:MAX_CONCEPT_SLIDES]:
        slide = add_content_slide(deck, concept["name"], footer=footer)
        add_text_box(slide, concept["explanation"], top=1.5, height=1.9, size=18)
        if concept["example"]:
            add_code_box(slide, concept["example"])

    if unit.tasks or unit.steps:
        items = unit.tasks or unit.steps
        add_bullet_slide(
            deck,
            "Your task this week",
            [f"{i}. {task}" for i, task in enumerate(items, 1)],
            notes=("Rules:\n" + "\n".join(f"- {rule}" for rule in unit.rules)) if unit.rules else "",
            footer=footer,
        )

    if unit.tests:
        add_bullet_slide(
            deck,
            "How you'll be graded",
            [
                "Your work is graded automatically by the unit's tests.",
                "Run the tests yourself before submitting.",
            ],
            notes="The tests check that:\n"
            + "\n".join(f"- {humanize_test(test)}" for test in unit.tests),
            footer=footer,
        )

    workflow = unit.sections.get("Starter workflow (GUI-only)") or unit.sections.get(
        "Starter workflow"
    )
    if workflow:
        bullets = [line.lstrip("-* ").strip() for line in workflow.splitlines() if line.strip()]
        add_bullet_slide(deck, "Workflow reminder", bullets, footer=footer)

    return deck_bytes(deck)


def _strip_md(text):
    """Inline markdown markers removed for Word body text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return text.replace("`", "").strip()


def _add_md_block(document, text):
    """Markdown-ish section text -> Word paragraphs (bullets become List
    Bullet style, everything else body text)."""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("-", "*")):
            document.add_paragraph(_strip_md(stripped.lstrip("-* ")), style="List Bullet")
        else:
            document.add_paragraph(_strip_md(stripped))


def _docx_bytes(document):
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _new_document():
    document = Document()
    document.styles["Normal"].font.size = DocxPt(11)
    return document


def build_lms_intro(unit, course_title):
    """Weekly LMS introduction as a professionally styled .docx."""
    document = _new_document()
    document.add_heading(f"Week {unit.week}: {unit.topic}", level=1)
    document.add_paragraph(
        f"Welcome to week {unit.week} of {course_title}! This week we tackle "
        f"{unit.topic}."
    )

    document.add_heading("What you'll learn", level=2)
    if unit.concepts:
        for concept in unit.concepts:
            document.add_paragraph(
                f"Understand {concept['name'].lower()}", style="List Bullet"
            )
    else:
        document.add_paragraph(f"The fundamentals of {unit.topic}", style="List Bullet")

    document.add_heading("This week's goal", level=2)
    target = unit.sections.get("Learning target", "").strip()
    document.add_paragraph(
        _strip_md(target)
        if target
        else f"Complete the {unit.slug} unit and make all of its automated tests pass."
    )

    document.add_heading("What to do", level=2)
    for step in (
        "Read this week's lecture slides.",
        f"Open {unit.path}/ in the project and read the instructions.",
        "Complete the starter code where marked, testing as you go.",
        "Commit your work and confirm the automated tests pass.",
    ):
        document.add_paragraph(step, style="List Number")

    document.add_heading("How you'll be graded", level=2)
    document.add_paragraph("Your submission is graded automatically and deterministically:")
    for test in unit.tests or []:
        document.add_paragraph(humanize_test(test), style="List Bullet")
    if not unit.tests:
        document.add_paragraph("The unit's automated checks pass", style="List Bullet")
    document.add_paragraph(
        "No placeholder/starter values remain in your code", style="List Bullet"
    )
    document.add_paragraph("Your code runs without syntax errors", style="List Bullet")

    document.add_paragraph(
        "Post in the discussion forum if you get stuck — struggling is part of "
        "learning, but staying stuck isn't."
    )
    return _docx_bytes(document)


def build_assignment_doc(unit):
    """Weekly assignment instructions as a professionally styled .docx."""
    document = _new_document()
    document.add_heading(f"Week {unit.week} Assignment — {unit.topic}", level=1)

    if unit.sections.get("Learning target"):
        document.add_heading("Goal", level=2)
        _add_md_block(document, unit.sections["Learning target"])

    if unit.tasks:
        document.add_heading("Your job", level=2)
        for task in unit.tasks:
            document.add_paragraph(_strip_md(task), style="List Number")
    elif unit.steps:
        document.add_heading("Steps", level=2)
        for step in unit.steps:
            document.add_paragraph(_strip_md(step), style="List Number")

    if unit.rules:
        document.add_heading("Rules your solution must follow", level=2)
        for rule in unit.rules:
            document.add_paragraph(_strip_md(rule), style="List Bullet")

    workflow = unit.sections.get("Starter workflow (GUI-only)") or unit.sections.get(
        "Starter workflow"
    )
    if workflow:
        document.add_heading("Workflow", level=2)
        _add_md_block(document, workflow)

    for heading, text in unit.sections.items():
        if heading.lower().startswith(("learning target", "starter workflow")):
            continue
        document.add_heading(_strip_md(heading), level=2)
        _add_md_block(document, text)

    if unit.test_file:
        document.add_heading("Check your work", level=2)
        document.add_paragraph(
            f"Run the automated tests in {unit.test_file}. Your submission is "
            "graded on these exact tests — if they pass for you, they pass for "
            "the grader."
        )
        for test in unit.tests:
            document.add_paragraph(humanize_test(test), style="List Bullet")

    return _docx_bytes(document)


def build_rubric(units):
    """One deterministic rubric for the whole course, consumable by a non-LLM
    grader. Every criterion is a mechanical check; weights sum to 100."""
    rubric_units = []
    for unit in units:
        criteria = []
        if unit.tests:
            criteria.append(
                {
                    "id": "tests_pass",
                    "type": "pytest",
                    "target": unit.test_file,
                    "tests": unit.tests,
                    "scoring": "weight * (passed_tests / total_tests)",
                    "weight": 60,
                }
            )
        if unit.placeholder_lines:
            criteria.append(
                {
                    "id": "placeholders_resolved",
                    "type": "forbidden_lines_absent",
                    "file": unit.code_file,
                    "forbidden_lines": unit.placeholder_lines,
                    "scoring": "weight * (1 - found_lines / total_forbidden_lines)",
                    "weight": 20,
                }
            )
        if unit.code_file:
            criteria.append(
                {
                    "id": "code_compiles",
                    "type": "python_compiles",
                    "file": unit.code_file,
                    "scoring": "weight if py_compile succeeds else 0",
                    "weight": 10,
                }
            )
        required = sorted(
            p for p in [unit.code_file, unit.test_file, f"{unit.path}/INSTRUCTIONS.md"] if p
        )
        criteria.append(
            {
                "id": "required_files_present",
                "type": "files_present",
                "required": required,
                "scoring": "weight * (present_files / required_files)",
                "weight": 10,
            }
        )
        # Renormalize so weights always sum to exactly 100.
        total = sum(c["weight"] for c in criteria)
        scaled = [int(c["weight"] * 100 / total) for c in criteria]
        scaled[0] += 100 - sum(scaled)
        for criterion, weight in zip(criteria, scaled):
            criterion["weight"] = weight

        rubric_units.append(
            {
                "id": unit.slug,
                "week": unit.week,
                "topic": unit.topic,
                "path": unit.path,
                "max_points": 100,
                "criteria": criteria,
            }
        )

    return {
        "version": 1,
        "scale": {"max_points_per_unit": 100},
        "grader_contract": {
            "pytest": "Run the listed tests against the submission with pytest; "
            "score = weight * passed/total.",
            "forbidden_lines_absent": "Strip whitespace from each submission line; "
            "score = weight * (1 - matches/total) where a match is any forbidden "
            "line appearing verbatim.",
            "python_compiles": "Byte-compile the file (py_compile); full weight "
            "on success, zero on SyntaxError.",
            "files_present": "score = weight * (present/required).",
        },
        "units": rubric_units,
    }


CSV_COLUMNS = [
    "unit_id",
    "week",
    "topic",
    "unit_path",
    "max_points",
    "criterion_id",
    "criterion_type",
    "weight",
    "scoring",
    "target",
    "details",
]
_DETAIL_SEPARATOR = " | "


def rubric_to_csv(rubric):
    """The rubric as one flat deterministic CSV: a GRADER_CONTRACT row per
    criterion type (how to evaluate it), then one row per unit criterion.
    Multi-value cells (test names, forbidden lines, required files) are
    joined with ' | '."""
    output = io.StringIO()
    writer = csv.writer(output, lineterminator="\n")
    writer.writerow(CSV_COLUMNS)
    for criterion_type in sorted(rubric["grader_contract"]):
        writer.writerow(
            [
                "GRADER_CONTRACT",
                "",
                "",
                "",
                "",
                "",
                criterion_type,
                "",
                "",
                "",
                rubric["grader_contract"][criterion_type],
            ]
        )
    for unit in rubric["units"]:
        for criterion in unit["criteria"]:
            if criterion["type"] == "pytest":
                target = criterion["target"]
                details = _DETAIL_SEPARATOR.join(criterion["tests"])
            elif criterion["type"] == "forbidden_lines_absent":
                target = criterion["file"]
                details = _DETAIL_SEPARATOR.join(criterion["forbidden_lines"])
            elif criterion["type"] == "python_compiles":
                target = criterion["file"]
                details = ""
            else:  # files_present
                target = ""
                details = _DETAIL_SEPARATOR.join(criterion["required"])
            writer.writerow(
                [
                    unit["id"],
                    unit["week"],
                    unit["topic"],
                    unit["path"],
                    unit["max_points"],
                    criterion["id"],
                    criterion["type"],
                    criterion["weight"],
                    criterion["scoring"],
                    target,
                    details,
                ]
            )
    return output.getvalue()


def build_manifest(units):
    """Bundle manifest as a .docx with a week/unit/topic table."""
    document = _new_document()
    document.add_heading("Course materials", level=1)
    document.add_paragraph(
        f"Generated deterministically from the uploaded project ({len(units)} units)."
    )

    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text = "Week", "Unit", "Topic"
    for unit in units:
        row = table.add_row().cells
        row[0].text = str(unit.week)
        row[1].text = unit.slug
        row[2].text = unit.topic

    document.add_heading("Contents", level=2)
    for line in (
        "lectures/ — one PPTX lecture per week",
        "lms/ — weekly LMS introduction posts (.docx)",
        "assignments/ — weekly assignment instructions (.docx)",
        "rubric.csv — deterministic rubric for a non-LLM grader "
        "(GRADER_CONTRACT rows document each criterion type)",
    ):
        document.add_paragraph(line, style="List Bullet")
    return _docx_bytes(document)


def _slugify(text):
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")[:40] or "unit"


def build_materials(zip_bytes, course_title="this course"):
    """The full course-materials zip (bytes) plus a summary dict."""
    units = parse_project(zip_bytes)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as bundle:
        for unit in units:
            stem = f"week-{unit.week:02d}-{_slugify(unit.topic)}"
            bundle.writestr(f"lectures/{stem}.pptx", build_lecture(unit))
            bundle.writestr(f"lms/{stem}.docx", build_lms_intro(unit, course_title))
            bundle.writestr(f"assignments/{stem}.docx", build_assignment_doc(unit))
        bundle.writestr("rubric.csv", rubric_to_csv(build_rubric(units)))
        bundle.writestr("MANIFEST.docx", build_manifest(units))

    summary = {
        "units": len(units),
        "weeks": [{"week": u.week, "unit": u.slug, "topic": u.topic} for u in units],
    }
    return output.getvalue(), summary
